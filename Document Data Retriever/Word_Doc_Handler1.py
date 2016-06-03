# -*- coding: utf-8 -*-
"""
Created on Thu May 26 12:37:22 2016

@author: mutabesham
"""
import win32com.client
import os
#import hashlib
import docx

import csv

class fileHandler_Error(Exception):
    """"
    Application error
    """
    def __init__(self, value):
        self.value = value
    def __str__(self):
        return repr(self.value)
        

#from extract_text import write_header, write_contents, write_path_names, get_files_processed, vectorize_document, files_processed
class fileHandler:
    """"
    Common class for all type of document wrapper, f.e: wordWrapper, ExcelWarpper..
    """
    #all variables 
    #global file_paths  # List which will store all of the full filepaths.
    #global filePath
    #global type_app
    
    def __init__(self, path, typOfapp = "word.Application"):
         self.type_app = typOfapp
        
    def __get_list_paths__(self,path):
        file_paths = [] #instantiate the list of the full filepaths
        if os.path.isdir(path): # if the directory path not provided       
            # Walk the tree.
            for root, directories, files in os.walk(path):
                for filename in files:
                    # Join the two strings in order to form the full filepath.
                    filepath = os.path.join(root, filename)
                    file_paths.append(filepath)  # Add it to the list.
            return file_paths
            #if typOfapp == "" we can check if the application type is the type we have implemented        
        elif os.path.isfile(path):
            file_paths.append(path)
            return file_paths # this case it contain only one filepath
        else :
            raise fileHandler_Error('Not a valide path')
            

    def upgrade_Doc_ToDocx(self,pathFolder):
        """
        this function takes folder/file path and save all word document 
        in doc version to new version docx in the same forlder and keep the doc version too
        """
        try:
            wordApp = win32com.client.Dispatch(self.type_app) # instatiate the application
            wordApp.Visible = 1 # make wordapp work in background
            full_file_paths = self.__get_list_paths__(pathFolder) 
            for this_path in full_file_paths: 
                if this_path.endswith(('.doc')) and not this_path.startswith(('~')):
                    tempDoc = wordApp.Documents.Add(Template= this_path) # get exactly the copy of the file
                    tempDoc.SaveAs(this_path.replace('.doc', '.docx'),FileFormat = 7 )
            tempDoc.Close()
            wordApp.Quit()
        except Exception as e:
            print(e.args)
        except fileHandler_Error as e:
            print(e.args)
        finally:
            
            return True
#    def get_file_ftp(self, destFolder)
            

class wordDocumentWrapper(fileHandler):
    
    FILE_SPLITER = "Profiel"
    
    def __init__(self, path, typOfapp = "word.Application"):
        super().__init__(path = path) # the same constractor as the parent class
    
    
    def __split_File__(self, path,destinationFolder):
        """
        this splite the cv in two version english and dutch
        """       
        try:
#            wordApp = win32com.client.Dispatch(self.type_app) # instatiate the application
#            wordApp.Visible = 1 # make wordapp work in background
            if os.path.isfile(path): # if file exist and not empty
#                myDoc = wordApp.Documents.Add(path)
                doc = docx.Document(path)
                file_current = docx.Document()
                english_file = docx.Document()
                isDutch_Present = False
                isEglish_Present = False
                English_dir = destinationFolder + "\\" + "EngishFolder"
                Dutch_dir = destinationFolder + "\\" + "DutchFolder"
                if not os.path.isdir(English_dir) and not os.path.isdir(Dutch_dir):
                        os.makedirs(destinationFolder + "\\" + "EngishFolder")
                        os.makedirs(destinationFolder + "\\" + "DutchFolder")
                #text = myDoc.ActiveDocument.Sections[0].Headers[win32.constants.wdHeaderFooterPrimary].Range.Text
                #print(text)
                for para in doc.paragraphs:
                    if para.text != "Profiel": # check if there is english 
                        file_current.add_paragraph(para.text) 
                        isEglish_Present = True
                    else:
                        isDutch_Present = True
                        english_file = file_current
                        file_current = docx.Document() # reset it for english version
                        file_current.add_paragraph("Profiel") # adding Profiel on top of the file 
                        
                if isDutch_Present and isEglish_Present:                  
                    new_path = English_dir+ "\\"+ os.path.basename(path).replace('.docx', 'English.docx')
                    english_file.save(new_path)
                    new_path = Dutch_dir+ "\\"+ os.path.basename(path).replace('.docx', 'Dutch.docx')
                    file_current.save(new_path)
                elif isEglish_Present and not isDutch_Present:
                    new_path = English_dir+ "\\"+ os.path.basename(path).replace('.docx', 'English.docx')
                    file_current.save(new_path)
                else:
                    new_path = Dutch_dir+ "\\" + os.path.basename(path).replace('.docx', 'Dutch.docx')
                    english_file.save(new_path)
                    help(english_file)
                
#                myDoc.Close()
#            wordApp.Quit()
        except Exception as e:
            print(e.args)
        except fileHandler_Error as e:
            print(e.args)
        finally:
            return True
    
    def split_Files(self,pathFolder, destinationFolder):
        """
        this function takes folder/file path and save all word document 
        in doc version to new version docx in the same forlder and keep the doc version too
        """
        try:
            wordApp = win32com.client.Dispatch(self.type_app) # instatiate the application
            wordApp.Visible = 0 # make wordapp work in background
            full_file_paths = self.__get_list_paths__(pathFolder) 
            for this_path in full_file_paths: 
                if this_path.endswith(('.doc')) and not this_path.startswith(('~')):
                    tempDoc = wordApp.Documents.Add(Template= this_path) # get exactly the copy of the file
                    new_path = this_path.replace('.doc', '.docx')
                    tempDoc.SaveAs(new_path )
                    self.__split_File__(new_path,destinationFolder) # this split a file and save in two versions(english and dutch)
                    tempDoc.Close()
                elif this_path.endswith(('.docx')):
                    self.__split_File__(this_path,destinationFolder) # this split a file and save in two versions(english and dutch)         
            wordApp.Quit()
        except Exception as e:
            print(e.args)
        except fileHandler_Error as e:
            print(e.args)
        finally:
            return True
    def __column_matcher__(self, text, dic):
        for k in dic:
            if k == text:
                return dic[k]
        return -1 # if not match
        
    def convert_to_CSV(self, path):
        
        CSV_DELIMITER = "|"
        HEADER_DIC = {'Profile':1, 'Industry Experience':2 , 'Skills & Competencies':3, 'Career Summary':4, 
                       'Career History':5, 'Career History':6 , 'Business Skills': 7, 'IT Skills':8 , 'Languages':9,
                       'Qualifications / Affiliations':10,'Skill Matrix ':11}
        SKILL_MATRIC_DIC ={'Methods':1,'Techniques':2,'Platforms':3,'Tools':4,'Programming environments & Databases':5,'Themes & Markets':6}
       
        CSV_FIELD_NAMES = ['Profile', 'Industry Experience' , 'Skills & Competencies', 'Career Summary',
                           'Career History', 'Career History' , 'Business Skills', 'IT Skills' , 'Languages',
                           'Qualifications / Affiliations''Methods','Techniques','Platforms','Tools',
                           'Programming environments & Databases','Themes & Markets']        
        DATA_ROW_LIST = []
        
        full_file_paths = self.__get_list_paths__(path) 
        
        with open(os.path.abspath(os.path.join(path_dest, os.pardir)) + "\\Data_csv.csv", 'w', newline='') as csvFile:                    
             csv_writer = csv.DictWriter(csvFile,fieldnames= CSV_FIELD_NAMES)
             csv_writer.writeheader()
        for this_path in full_file_paths:
            if os.path.isfile(this_path) and this_path.endswith(('.docx')) and not this_path.startswith(('~')): # if file exist and not empty
                doc = docx.Document(this_path)
         
                key_founded_list = []
                current_key = ""
                row_data_dic= {}
                row_temp = ""
                for key in HEADER_DIC.keys():
                    for para in doc.paragraphs:
                        if para.text == key:
                            print(para.text)
                            key_founded_list.append(key)
                            if row_temp != "":
                                row_data_dic[current_key] = row_temp
                                row_temp= ""
                            current_key = key
#                            continue # go to the next paragraph
                        elif(current_key in key_founded_list ):
                            row_temp.join(" "+ para.text)
                            row_data_dic[current_key] = str(row_data_dic.get(current_key)) + " " + row_temp
                            
                    if len(row_data_dic) > 0:
                        DATA_ROW_LIST.append(row_data_dic)

                            
#            """this code is for looping in the skill matric """
#            if len(doc.tables) > 0 :
#                for table in doc.tables:
#                    for row in table.rows:
#                        for cell in row.cells:  
#                            match = __column_matcher__(cell.text,HEADER_Dic)
#                            if math == 11: # if matches the skill matrics
#                               print(cell.text)
       
                
        for row in DATA_ROW_LIST:
            csv_writer.writerow(row)
                       
        
    
        
        
 ############################## code for client app goes here ######################################
        
word_doc = wordDocumentWrapper('C:\\Users/mutabesham\\Documents')
path = 'C:\\Users\\mutabesham\\Documents\\CVs'
path_dest = 'C:\\Users\\mutabesham\\Documents\\CVs\\plited cvs'

'''Decomment line below if the folder contain old version of words documnets'''
#word_doc.upgrade_Doc_ToDocx(path)

if word_doc.split_Files(path, path_dest):
     print("Files splited!")
#word_doc.convert_to_CSV(path)
